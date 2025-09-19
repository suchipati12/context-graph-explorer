"""
File parsing utilities for Context Graph Explorer
Supports PDF, DOCX, and Markdown file parsing
"""

import os
import tempfile
from typing import Dict, Any, Optional
import streamlit as st

# PDF parsing
try:
    from pypdf import PdfReader
except ImportError:
    st.error("pypdf not installed. Please install it using: pip install pypdf")

# DOCX parsing
try:
    from docx import Document
except ImportError:
    st.error("python-docx not installed. Please install it using: pip install python-docx")


class DocumentParser:
    """Main document parser class supporting multiple file formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']
    
    def parse_document(self, uploaded_file) -> Dict[str, Any]:
        """
        Parse uploaded document and extract text content
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dict containing parsed content and metadata
        """
        if uploaded_file is None:
            return {"error": "No file uploaded"}
        
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension not in self.supported_formats:
            return {"error": f"Unsupported file format: {file_extension}"}
        
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            # Parse based on file type
            if file_extension == '.pdf':
                result = self._parse_pdf(tmp_file_path, uploaded_file.name)
            elif file_extension == '.docx':
                result = self._parse_docx(tmp_file_path, uploaded_file.name)
            elif file_extension in ['.txt', '.md']:
                result = self._parse_text(tmp_file_path, uploaded_file.name)
            else:
                result = {"error": f"Parser not implemented for {file_extension}"}
            
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
            return result
            
        except Exception as e:
            return {"error": f"Error parsing document: {str(e)}"}
    
    def _parse_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse PDF document"""
        try:
            reader = PdfReader(file_path)
            text_content = ""
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text_content += f"\n--- Page {page_num + 1} ---\n"
                text_content += page_text
            
            return {
                "filename": filename,
                "file_type": "PDF",
                "num_pages": len(reader.pages),
                "text_content": text_content.strip(),
                "metadata": {
                    "pages": len(reader.pages),
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            return {"error": f"Error parsing PDF: {str(e)}"}
    
    def _parse_docx(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            return {
                "filename": filename,
                "file_type": "DOCX",
                "num_paragraphs": len(doc.paragraphs),
                "text_content": text_content.strip(),
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            return {"error": f"Error parsing DOCX: {str(e)}"}
    
    def _parse_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse text/markdown document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            lines = text_content.split('\n')
            
            return {
                "filename": filename,
                "file_type": "TEXT/MARKDOWN",
                "num_lines": len(lines),
                "text_content": text_content.strip(),
                "metadata": {
                    "lines": len(lines),
                    "file_size": os.path.getsize(file_path)
                }
            }
        except Exception as e:
            return {"error": f"Error parsing text file: {str(e)}"}
    
    def get_text_preview(self, text_content: str, max_chars: int = 500) -> str:
        """Get a preview of the text content"""
        if len(text_content) <= max_chars:
            return text_content
        return text_content[:max_chars] + "..."


def validate_file_upload(uploaded_file) -> Optional[str]:
    """Validate uploaded file and return error message if invalid"""
    if uploaded_file is None:
        return "Please upload a file"
    
    max_size = 10 * 1024 * 1024  # 10MB
    if uploaded_file.size > max_size:
        return f"File size too large. Maximum allowed size is {max_size // (1024*1024)}MB"
    
    allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    
    if file_extension not in allowed_extensions:
        return f"Unsupported file type. Allowed types: {', '.join(allowed_extensions)}"
    
    return None